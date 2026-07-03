#!/usr/bin/env python3
"""Generate user manual for tax-policy-search project as Word document."""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"D:\Desktop\桌面迁移\tax-policy-search\财税政策搜索引擎_用户手册.docx"

doc = Document()

# ═══════════════════════════════════════════════════════════════
#  STYLES SETUP
# ═══════════════════════════════════════════════════════════════
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
# Set East Asian font
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Microsoft YaHei'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    h.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════════════════════════
#  COVER / TITLE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('财税政策实时搜索引擎')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Tax Policy Real-Time Search Engine\n用户手册')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('版本：v1.2.0\n').font.size = Pt(10)
meta.add_run('更新日期：2026-07-03\n').font.size = Pt(10)
meta.add_run('GitHub：https://github.com/huang-doudou567/tax-policy-search\n').font.size = Pt(10)
meta.add_run('Demo：https://huang-doudou567.github.io/tax-policy-search/').font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '一、项目概述',
    '二、系统架构',
    '三、功能模块详解',
    '四、数据源说明',
    '五、使用方法',
    '六、版本迭代记录',
    '七、项目目录结构',
    '八、技术栈',
    '九、常见问题（FAQ）',
    '十、免责声明',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  一、项目概述
# ═══════════════════════════════════════════════════════════════
doc.add_heading('一、项目概述', level=1)

doc.add_heading('1.1 项目定位', level=2)
doc.add_paragraph(
    '财税政策实时搜索引擎是一个面向企业财务人员、税务师、个体工商户及法务合规人员的中国税收政策法规查询与解读服务平台。'
    '系统通过连接全国人民代表大会（NPC）国家法律法规数据库、国家税务总局（chinatax.gov.cn）等权威数据源，'
    '提供 12 大税种的最新政策法规查询、申报指导、优惠资格判定与合规风险提示。'
)

doc.add_heading('1.2 核心原则', level=2)
principles = [
    '先搜索，后回答 — 永远不凭训练数据直接给出税收政策答案',
    '实时查询 — 每个回答必须包含实时查询时间戳',
    '时效性标注 — 明确区分现行有效 / 已修改 / 已废止',
    '免责声明 — 所有回答强制附加免责声明',
]
for p in principles:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('1.3 目标用户', level=2)
add_table(doc, ['用户群', '典型场景'], [
    ['企业主 / 财务人员', '"小微企业增值税有什么优惠？""研发费用加计扣除比例是多少？"'],
    ['个体工商户', '"小规模纳税人开票限额？""个体户需要汇算清缴吗？"'],
    ['代理记账 / 税务师', '"最新增值税优惠政策汇总""金税四期风险指标有哪些？"'],
    ['法务 / 合规人员', '"关联交易转让定价规定""跨境税务合规要求"'],
], col_widths=[4, 12])

doc.add_heading('1.4 项目规模', level=2)
add_table(doc, ['指标', '数值'], [
    ['项目文件总数', '69 个'],
    ['总代码行数', '约 6,700 行（含 Python + HTML + JS + MD）'],
    ['Python 脚本', '6 个核心模块，约 2,600 行'],
    ['前端页面', '3 个 HTML（含 docs/ + frontend/）'],
    ['测试用例', '11 项端到端测试'],
    ['数据文件', '21 个 JSON 搜索索引'],
    ['Git 提交', '25 次'],
], col_widths=[5, 11])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  二、系统架构
# ═══════════════════════════════════════════════════════════════
doc.add_heading('二、系统架构', level=1)

doc.add_heading('2.1 总体架构图', level=2)
doc.add_paragraph(
    '系统采用"双模架构"设计，同时支持本地完整版和 GitHub Pages 静态 Demo 版。'
)
doc.add_paragraph(
    '【本地完整版架构】\n'
    '  用户浏览器 → Flask API Server (tax_server.py)\n'
    '    ├── POST /api/search        → NPC 国家法规库 API（flk.npc.gov.cn）\n'
    '    ├── GET  /api/text/<id>      → DOCX 下载解析（Python stdlib）\n'
    '    ├── GET  /api/interpretations → Bing 搜索 chinatax / mof / gov.cn\n'
    '    ├── GET  /api/web-related    → 实务公众号 + 网页搜索\n'
    '    └── GET  /api/ai-interpret   → Claude Code CLI 生成解读\n'
    '\n'
    '【GitHub Pages Demo 架构】\n'
    '  用户浏览器 → GitHub Pages (docs/index.html)\n'
    '    ├── docs/data/index.json      → 预置搜索索引（207KB）\n'
    '    ├── GitHub Actions 每日自动更新 → 从 NPC API 拉取最新数据\n'
    '    └── 纯客户端 JavaScript 搜索 → 毫秒级响应'
)

doc.add_heading('2.2 模块层次', level=2)
add_table(doc, ['层级', '文件', '职责'],
[
    ['前端层', 'frontend/index.html (430行)', '搜索界面 + 智能引导 + 法规弹窗 4-Tab'],
    ['API 服务', 'scripts/tax_server.py (710行)', 'Flask 路由 + 法规原文 + 解读搜索 + AI 解读'],
    ['核心搜索', 'scripts/tax_search.py (492行)', '18 税种映射 + NPC API 调用 + 意图识别 + 缓存'],
    ['法规详情', 'scripts/tax_detail.py (240行)', 'NPC 详情 API + DOCX 下载 + 全文解析（zipfile+ElementTree）'],
    ['Web 搜索', 'scripts/tax_web_search.py (240行)', 'chinatax WAS5 搜索引擎 + Bing site: + Baidu 兜底'],
    ['多源聚合', 'scripts/tax_aggregator.py (274行)', 'ThreadPoolExecutor 三源并发 + Jaccard 去重 + 权威度排序'],
    ['输出格式化', 'scripts/tax_formatter.py (198行)', '四段式 Markdown + 时效性徽章 + 免责声明'],
    ['数据采集', '.github/scripts/fetch_tax_data.py', 'GitHub Actions 每日从 NPC API 获取 20 个搜索集'],
    ['静态 Demo', 'docs/index.html (434行)', '纯前端搜索 + 智能引导 + 6 项筛选 + 关键词高亮'],
    ['部署1', 'deploy-vercel/', 'Vercel Serverless Python 部署方案'],
    ['部署2', 'cloudflare-deploy/', 'Cloudflare Pages + Workers 部署方案'],
], col_widths=[3, 5, 8])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  三、功能模块详解
# ═══════════════════════════════════════════════════════════════
doc.add_heading('三、功能模块详解', level=1)

doc.add_heading('3.1 实时法规检索', level=2)
doc.add_paragraph(
    '核心搜索模块连接 NPC 国家法律法规数据库 API，支持标题/全文、精确/模糊搜索，'
    '并提供时效性过滤、日期范围筛选、排序和分页功能。'
)
add_table(doc, ['参数', '选项', '说明'],
[
    ['搜索范围 (scope)', 'title / fulltext', '标题搜索优先，无结果自动降级全文搜索'],
    ['匹配方式 (search_type)', 'exact / fuzzy', '精确匹配用于已知法规名，模糊匹配用于宽泛主题'],
    ['时效性 (status)', '1=已废止, 2=已修改, 3=现行有效, 4=尚未生效', '默认仅查现行有效 (sxx=3)'],
    ['排序 (sort)', 'relevance / date', '日期排序支持按公布日期降序查看最新政策'],
    ['分页', '1-50 条/页', 'NPC API 最大支持 100 条/页'],
    ['日期范围', '起始—截止', '支持按公布日期范围筛选'],
], col_widths=[4, 5, 7])

doc.add_heading('3.2 意图识别引擎', level=2)
doc.add_paragraph('系统自动分析用户输入，将查询分为 5 种意图类型，匹配最优搜索策略：')
add_table(doc, ['意图', '触发信号', '搜索策略'],
[
    ['📖 政策查询', '多少、比例、最新、规定、优惠', '标题精确 → 全文模糊兜底'],
    ['📋 申报指导', '申报、汇算清缴、截止、流程', '全文搜索 + chinatax 操作指南补充'],
    ['⚠️ 合规风险', '风险、金税、会被查、预警', '全文搜索风险关键词 + 指标框架参考'],
    ['🎁 资格判定', '符合条件、能不能享受、是否适用', '精确搜索政策名 + 全文条件关键词'],
    ['🧾 发票处理', '开票、红冲、遗失、抵扣认证', '"发票管理办法"精确 + 关键词全文'],
], col_widths=[3, 5, 8])

doc.add_heading('3.3 18 税种关键词映射', level=2)
doc.add_paragraph('系统内置税种别名映射，将用户口语自动翻译为法律术语，提高搜索命中率。')
add_table(doc, ['用户说', '搜索关键词', '优先法规'],
[
    ['增值税/VAT', '增值税', '中华人民共和国增值税法'],
    ['企业所得税', '企业所得税', '中华人民共和国企业所得税法'],
    ['个税/个人所得税', '个人所得税', '中华人民共和国个人所得税法'],
    ['消费税', '消费税', '中华人民共和国消费税法'],
    ['关税/海关税', '关税', '中华人民共和国关税法'],
    ['房产税/房地产税', '房产税', '中华人民共和国房产税暂行条例'],
    ['土地增值税/土增税', '土地增值税', '中华人民共和国土地增值税暂行条例'],
    ['契税', '契税', '中华人民共和国契税法'],
    ['印花税', '印花税', '中华人民共和国印花税法'],
    ['城建税', '城市维护建设税', '中华人民共和国城市维护建设税法'],
    ['环保税', '环境保护税', '中华人民共和国环境保护税法'],
    ['发票问题', '发票管理办法', '中华人民共和国发票管理办法'],
], col_widths=[4, 4, 8])

doc.add_heading('3.4 智能引导面板（4 步向导）', level=2)
doc.add_paragraph(
    '帮助不确定查什么的用户精准定位政策需求，4 个步骤自动拼接搜索关键词并设置筛选条件：\n'
    '\n'
    'Step 1: 选择身份 → 🏢企业 / 👤个人个体户 / 📋代理记账税务师 / 💬其他\n'
    'Step 2: 选择税种 → 12 个税种卡片（增值税、企业所得税、个人所得税等）\n'
    'Step 3: 选择意图 → 📖查政策 / 🎁优惠资格 / ⚠️风险自查 / 📋申报流程 / 🧾发票问题 / 🔢计算标准\n'
    'Step 4: 补充条件 → 纳税人类型(小规模/一般/小微/高新) + 时效(最新/2024以来) + 省份 + 自定义关键词\n'
    '\n'
    '点击「🔍 生成搜索」后自动拼接关键词并触发搜索。'
)

doc.add_heading('3.5 法规弹窗（本地完整版）', level=2)
doc.add_paragraph('点击搜索结果卡片弹出 4-Tab 法规弹窗：')
add_table(doc, ['Tab', '功能', '数据来源'],
[
    ['📖 法规原文', '全文展示 + 搜索关键词黄色高亮 + 章节/法条自动分类', 'NPC API → DOCX 实时下载解析 (Python stdlib)'],
    ['🔍 官方解读', '多源搜索官方政策解读、答记者问、立法说明', 'Bing 搜索 chinatax/mof/gov.cn'],
    ['🤖 AI 解读', 'Claude 用通俗语言解读法规（适用主体/核心要点/注意事项）', 'Claude Code CLI 子进程'],
    ['🌐 相关网页', '公众号实务解读（小颖言税/税海涛声/会计网等）+ 网络补充', 'Bing + Baidu 双引擎'],
], col_widths=[3, 6, 7])

doc.add_heading('3.6 高级筛选栏', level=2)
add_table(doc, ['筛选', '选项', '说明'],
[
    ['范围', '标题 / 全文', '指定搜索范围'],
    ['匹配', '模糊 / 精确', '精确仅匹配标题含完整关键词'],
    ['时效', '现行有效（默认）/ 全部', '筛选法规状态'],
    ['排序', '相关度 / 日期↓', '排序方式'],
    ['省份', '全国 + 36 个省市', '按发布机关省份筛选地方性法规'],
    ['日期', '起—止范围', '按公布日期区间筛选'],
], col_widths=[3, 5, 8])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  四、数据源说明
# ═══════════════════════════════════════════════════════════════
doc.add_heading('四、数据源说明', level=1)

doc.add_heading('4.1 权威数据源', level=2)
add_table(doc, ['数据源', '覆盖范围', '权威度', '访问方式'],
[
    ['NPC 国家法规库\n(flk.npc.gov.cn)', '法律、行政法规、地方法规、司法解释', '⭐⭐⭐⭐⭐', 'API 实时'],
    ['国家税务总局\n(chinatax.gov.cn)', '部门规章、公告、政策解读、操作指南', '⭐⭐⭐⭐', 'WebFetch'],
    ['财政部\n(mof.gov.cn)', '财政部公告、财税联合发文', '⭐⭐⭐⭐', 'Bing site:'],
    ['中国政府网\n(gov.cn)', '国务院政策发布、答记者问', '⭐⭐⭐⭐⭐', 'Bing site:'],
    ['税务法规库\n(fgk.chinatax.gov.cn)', '税务法规库结构化内容', '⭐⭐⭐⭐', 'Bing site:'],
], col_widths=[4, 6, 2, 4])

doc.add_heading('4.2 实务解读来源', level=2)
doc.add_paragraph('为弥补官方解读偏正式的不足，系统引入了以下实务公众号作为补充参考源：')
add_table(doc, ['来源名称', '类型', '说明'],
[
    ['小颖言税', '微信公众号', '财税实务解读与案例分析'],
    ['税海涛声', '微信公众号', '税收政策深度剖析'],
    ['会计网', '微信公众号', '会计与税务实务操作指南'],
    ['税小课', '微信公众号', '税务知识培训与解读'],
    ['朴税', '微信公众号', '税务合规与风险分析'],
], col_widths=[4, 3, 9])

doc.add_heading('4.3 内容过滤机制', level=2)
doc.add_paragraph(
    '系统内置了内容过滤机制，在搜索实务解读和网络来源时，会自动拦截以下类型的非财税内容：\n'
    '\n'
    '• 屏蔽域名：游戏类(game/4399/7k7k)、视频类(bilibili/douyin/kuaishou)、社交类(weibo/tieba)\n'
    '• 屏蔽关键词：游戏、小游戏、手游、娱乐、八卦、彩票、博彩、小说、漫画\n'
    '\n'
    '确保搜索结果与税收政策相关。'
)

doc.add_heading('4.4 缓存策略', level=2)
add_table(doc, ['数据类型', '缓存时间', '说明'],
[
    ['搜索结果', '5 分钟', '政策随时更新，短缓存确保新鲜度'],
    ['详情元数据', '1 小时', '法规元信息变动频率低'],
    ['DOCX 文件', '24 小时', '法规全文极少变动'],
    ['默认行为', '不使用缓存', '每次必须实时查询'],
], col_widths=[4, 4, 8])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  五、使用方法
# ═══════════════════════════════════════════════════════════════
doc.add_heading('五、使用方法', level=1)

doc.add_heading('5.1 在线 Demo（推荐）', level=2)
doc.add_paragraph(
    '直接访问 GitHub Pages 永久地址即可使用：\n'
    '\n'
    '🔗 https://huang-doudou567.github.io/tax-policy-search/\n'
    '\n'
    '功能：12 税种搜索 + 6 项筛选 + 智能引导 + 关键词高亮 + 每日自动更新数据\n'
    '无需安装任何软件，浏览器打开即用。'
)

doc.add_heading('5.2 本地运行完整版', level=2)
doc.add_paragraph('本地完整版包含法规原文下载、官方解读搜索、AI 解读等全部功能。')
doc.add_paragraph(
    '步骤 1: 克隆项目\n'
    '  git clone https://github.com/huang-doudou567/tax-policy-search.git\n'
    '  cd tax-policy-search\n'
    '\n'
    '步骤 2: 安装依赖\n'
    '  pip install flask requests urllib3\n'
    '\n'
    '步骤 3: 启动服务\n'
    '  python scripts/tax_server.py\n'
    '\n'
    '步骤 4: 打开浏览器访问\n'
    '  http://localhost:5080'
)

doc.add_heading('5.3 命令行搜索', level=2)
doc.add_paragraph(
    '除了 Web 界面，也可以通过命令行直接搜索：\n'
    '\n'
    '# 政策查询\n'
    'python scripts/tax_search.py "增值税" --status 3 --size 20\n'
    '\n'
    '# 全文搜索 + 按日期排序\n'
    'python scripts/tax_search.py "加计扣除" --scope fulltext --sort date\n'
    '\n'
    '# 精确查找\n'
    'python scripts/tax_search.py "中华人民共和国企业所得税法" --exact\n'
    '\n'
    '# 获取法规详情\n'
    'python scripts/tax_detail.py --info <法规ID>\n'
    '\n'
    '# 多源聚合搜索\n'
    'python scripts/tax_aggregator.py "小微企业优惠" --size 10\n'
    '\n'
    '# 搜索官方解读\n'
    'python scripts/tax_web_search.py "增值税" --size 10'
)

doc.add_heading('5.4 Claude Code Skill 模式', level=2)
doc.add_paragraph(
    '项目包含 SKILL.md 技能定义文件，可安装为 Claude Code Skill。\n'
    '在 Claude Code 对话中直接提问：\n'
    '\n'
    '"小微企业增值税有什么优惠？"\n'
    '"研发费用加计扣除比例是多少？"\n'
    '"金税四期有哪些风险指标？"\n'
    '\n'
    'Claude 将自动激活 Search Skill，实时搜索 NPC 法规库后回答。\n'
    '触发词包括：增值税、所得税、税务、发票、申报、优惠、风险预警、金税四期等。'
)

doc.add_heading('5.5 部署方案', level=2)
doc.add_paragraph(
    '项目提供 3 种部署方式：\n'
    '\n'
    '① GitHub Pages（已部署）\n'
    '  优点：永久在线，零成本，无需服务器\n'
    '  配置：Settings → Pages → Source: master branch, Folder: /docs\n'
    '\n'
    '② Vercel Serverless\n'
    '  优点：支持 Python Flask，实时 API 连接 NPC\n'
    '  部署：cd deploy-vercel && vercel deploy --prod\n'
    '  注意：中国大陆可能需要自定义域名\n'
    '\n'
    '③ Cloudflare Pages + Workers\n'
    '  优点：全球 CDN，Workers 国内可访问\n'
    '  部署：cd cloudflare-deploy && wrangler pages deploy .'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  六、版本迭代记录
# ═══════════════════════════════════════════════════════════════
doc.add_heading('六、版本迭代记录', level=1)

add_table(doc, ['版本', '日期', '说明'],
[
    ['v1.0.0', '2026-06-29', '初始版本：NPC API 搜索、DOCX 解析、Flask 服务、SKILL.md、前端 Demo'],
    ['v1.0.1', '2026-06-29', '代码审查修复：提取共享 DOCX 解析器、移除冗余分省税种分支、模块级正则'],
    ['v1.0.2', '2026-07-01', '三段式 CHECKPOINT 标记（确认/风险/兜底）+ SKILL.md 优化'],
    ['v1.1.0', '2026-07-02', '搜索结果高亮修复 + 引入 5 家实务公众号（小颖言税/税海涛声/会计网/税小课/朴税）+ Vercel/Cloudflare 部署'],
    ['v1.1.1', '2026-07-02', '添加 fetch 自动重试机制 + 空响应兜底，解决隧道不稳定查询失败'],
    ['v1.1.2', '2026-07-02', '修复 doSearch 函数头覆盖导致 JS 花括号不匹配 → 页面完全无响应'],
    ['v1.1.3', '2026-07-02', '严格分离官方解读与相关网页 + 内容垃圾过滤 + 公众号归属修正'],
    ['v1.1.4', '2026-07-02', '添加隧道守护脚本 + 多种 API 探测方案'],
    ['v1.2.0', '2026-07-03', '重磅：GitHub Pages 纯前端搜素引擎 + GitHub Actions 每日自动更新数据'],
    ['v1.2.1', '2026-07-03', '恢复智能引导面板 + 高级筛选栏（范围/匹配/时效/排序/省份/日期）'],
    ['v1.2.2', '2026-07-03', '省份筛选接入搜索逻辑 + 智能引导 Step4 新增省份选择 + 修复 guidePick 崩溃'],
    ['v1.2.3', '2026-07-03', '最终确认纯静态方案为公开 Demo + 回退 iframe 方案'],
], col_widths=[2, 3, 11])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  七、项目目录结构
# ═══════════════════════════════════════════════════════════════
doc.add_heading('七、项目目录结构', level=1)

doc.add_heading('7.1 核心目录', level=2)
add_table(doc, ['路径', '说明'],
[
    ['tax-policy-search/', '项目根目录'],
    ['├── SKILL.md', 'Claude Code Skill 定义文件（AI Agent 操作手册）'],
    ['├── README.md', '项目介绍（含架构图、API 文档、使用示例）'],
    ['├── ANALYSIS.md', '架构分析报告'],
    ['├── COMPLETE_PROJECT_PLAN.md', '完整项目规划'],
    ['├── requirements.txt', 'Python 依赖（flask, requests, urllib3）'],
    ['├── .gitignore', 'Git 忽略规则'],
    ['│'],
    ['├── scripts/', 'Python 核心脚本'],
    ['│   ├── tax_server.py', 'Flask API 服务（6 个 REST 端点 + AI 解读）'],
    ['│   ├── tax_search.py', 'NPC API 核心搜索（18 税种映射 + 意图识别 + 缓存）'],
    ['│   ├── tax_detail.py', '法规详情与下载（DOCX ZIP+XML 解析）'],
    ['│   ├── tax_web_search.py', 'chinatax.gov.cn 站内搜索（WAS5 + Bing + Baidu）'],
    ['│   ├── tax_aggregator.py', '多源聚合（并发 + Jaccard 去重 + 权威度排序）'],
    ['│   ├── tax_formatter.py', '结构化 Markdown 输出'],
    ['│   └── tunnel_daemon.py', '隧道守护脚本（自动重连）'],
    ['│'],
    ['├── frontend/', 'Web 前端（本地完整版）'],
    ['│   └── index.html', '搜索界面 + 智能引导 + 4-Tab 法规弹窗'],
    ['│'],
    ['├── docs/', 'GitHub Pages 静态 Demo'],
    ['│   ├── index.html', '纯前端搜索引擎 + 智能引导 + 6 项筛选'],
    ['│   └── data/', '预置搜索数据（21 个 JSON，~207KB）'],
    ['│'],
    ['├── references/', '参考文档'],
    ['│   ├── tax_categories.md', '18 税种 × API 参数映射'],
    ['│   ├── search_strategies.md', '5 意图 × 搜索策略交叉'],
    ['│   └── tax_risk_framework.md', '金税四期 200+ 指标搜索提示'],
    ['│'],
    ['├── tests/', '测试'],
    ['│   └── test_tax_search.py', '11 项端到端测试'],
    ['│'],
    ['├── .github/', 'GitHub Actions 工作流'],
    ['│   ├── workflows/update-data.yml', '每日自动更新数据'],
    ['│   └── scripts/fetch_tax_data.py', '从 NPC API 采集 20 个搜索集'],
    ['│'],
    ['├── deploy-vercel/', 'Vercel 部署方案（9 files）'],
    ['└── cloudflare-deploy/', 'Cloudflare Workers 部署方案（9 files）'],
], col_widths=[8, 8])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  八、技术栈
# ═══════════════════════════════════════════════════════════════
doc.add_heading('八、技术栈', level=1)

add_table(doc, ['层级', '技术', '用途'],
[
    ['后端框架', 'Flask 3.1', 'REST API 服务'],
    ['HTTP 客户端', 'requests + urllib3', 'NPC API 调用、Bing 搜索、DOCX 下载'],
    ['文档解析', 'Python stdlib (zipfile + ElementTree)', 'DOCX → 纯文本提取（零依赖）'],
    ['并发', 'concurrent.futures (ThreadPoolExecutor)', '多源并行查询（NPC + chinatax + AnySearch）'],
    ['前端', '纯 HTML/CSS/JS', '搜索界面 + 引导面板 + 法规弹窗（零框架）'],
    ['AI 解读', 'Claude Code CLI (subprocess)', '法规通俗化解读生成'],
    ['测试', 'Python unittest 模式', '11 项端到端测试（22/22 通过）'],
    ['部署 1', 'GitHub Pages', '静态 Demo 部署（docs/ 目录）'],
    ['部署 2', 'Vercel + @vercel/python', 'Serverless Flask 部署'],
    ['部署 3', 'Cloudflare Pages + Workers', 'Workers JS 版本部署'],
    ['CI/CD', 'GitHub Actions', '每日自动从 NPC API 更新搜索数据'],
    ['依赖', 'flask, requests, urllib3', '仅 3 个 pip 包'],
], col_widths=[4, 6, 6])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  九、常见问题
# ═══════════════════════════════════════════════════════════════
doc.add_heading('九、常见问题（FAQ）', level=1)

faqs = [
    ('Q: Demo 页面与本地完整版有什么区别？',
     'A: Demo 页面（GitHub Pages）使用预置数据实现毫秒级搜索，包含全部税种搜索、筛选、智能引导功能，'
     '数据每日自动更新。本地完整版额外提供法规原文 DOCX 下载、Bing 搜索官方解读、AI 生成解读、'
     '公众号实务解读的实时搜索功能。'),

    ('Q: 搜索数据多久更新一次？',
     'A: GitHub Actions 每日凌晨 2:00（UTC）自动运行，从 NPC 国家法律法规数据库 API 拉取最新数据，'
     '覆盖 12 个税种和 8 个热门专题，共约 200KB 搜索索引。'),

    ('Q: 免费 SSH 隧道为什么不稳定？',
     'A: serveo.net 是免费的临时隧道服务，在连接空闲或网络波动时会自动断开，无法保证 7×24 在线。'
     '这也是最终选择纯静态 GitHub Pages 方案的原因——完全不依赖第三方运行环境。'),

    ('Q: 如何贡献或反馈问题？',
     'A: 欢迎通过 GitHub Issues 提交反馈：https://github.com/huang-doudou567/tax-policy-search/issues'),

    ('Q: 项目可以商用吗？',
     'A: 本项目仅供税务政策查询参考，不构成税务建议。所有税收政策信息以国家税务总局官网、财政部官网、'
     '主管税务机关的最新公告为准。涉及具体税务处理，请咨询专业税务人员。'),

    ('Q: 技术支持联系方式？',
     'A: GitHub 项目页或邮件联系。税务咨询请拨打 12366 纳税服务热线。'),
]

for q, a in faqs:
    doc.add_heading(q, level=2)
    doc.add_paragraph(a)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  十、免责声明
# ═══════════════════════════════════════════════════════════════
doc.add_heading('十、免责声明', level=1)

doc.add_paragraph(
    '本项目（财税政策实时搜索引擎）仅供税务政策查询参考，不构成税务建议、法律意见或专业咨询。\n'
    '\n'
    '所有税收政策信息以以下官方来源为准：\n'
    '  • 国家税务总局官网：https://www.chinatax.gov.cn\n'
    '  • 财政部官网：https://www.mof.gov.cn\n'
    '  • 国家法律法规数据库：https://flk.npc.gov.cn\n'
    '  • 纳税服务热线：12366\n'
    '\n'
    '涉及具体税务处理、申报操作和税务筹划，请咨询专业税务人员或主管税务机关。\n'
    '\n'
    '项目中的 AI 解读功能使用 Claude 生成，仅供参考，以官方政策文件和主管税务机关解释为准。\n'
    '搜素引擎返回的法规结果可能存在延迟，建议重要决策前通过 NPC 官方网站核实最新状态。\n'
    '\n'
    '项目遵循开源协议，使用方应自行承担使用风险。'
)

# ═══════════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f'✅ 用户手册已生成: {OUTPUT}')
print(f'   文件大小: {os.path.getsize(OUTPUT):,} bytes')
